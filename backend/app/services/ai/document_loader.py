"""Shared utility: extract plain text from .md / .txt / .pdf / .docx files."""
from pathlib import Path


_SUPPORTED = {".md", ".txt", ".pdf", ".docx"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix not in _SUPPORTED:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {sorted(_SUPPORTED)}")

    raw = path.read_bytes()

    if suffix in {".md", ".txt"}:
        return raw.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        import pypdf, io
        reader = pypdf.PdfReader(io.BytesIO(raw))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        import docx, io
        doc = docx.Document(io.BytesIO(raw))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"Unhandled suffix '{suffix}'")


def extract_text_bytes(content: bytes, suffix: str, filename: str = "") -> str:
    """Same as extract_text but operates on raw bytes (for upload endpoints)."""
    suffix = suffix.lower()
    if suffix not in _SUPPORTED:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {sorted(_SUPPORTED)}")

    if suffix in {".md", ".txt"}:
        return content.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        import pypdf, io
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        import docx, io
        doc = docx.Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"Unhandled suffix '{suffix}'")
